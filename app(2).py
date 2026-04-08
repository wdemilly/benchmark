import io
import json
import zipfile
from dataclasses import dataclass, asdict
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Tuple

import pandas as pd
import streamlit as st
from docx import Document

try:
    import anthropic  # type: ignore
except Exception:
    anthropic = None

APP_TITLE = "Author-Immersion Chapter Writer"
DATA_DIR = Path("author_runs")
DATA_DIR.mkdir(exist_ok=True)
OUTPUTS_DIR = DATA_DIR / "outputs"
OUTPUTS_DIR.mkdir(exist_ok=True)
CSV_PATH = DATA_DIR / "runs.csv"

DEFAULT_SYSTEM = (
    "You are writing fiction, not explaining it. Remain fully inside the requested chapter. "
    "Do not add commentary before or after the chapter. Return plain text only."
)

DEFAULT_USER_PROMPT = (
    "You are not Claude. You are the author of the combined source texts document.\n\n"
    "You wrote all of the passages. The character profiles are your notes. The outline is your plan for this chapter.\n\n"
    "Read the documents from beginning to end. Don't just sample them. Write the chapter now from the outline, constructing each sentence the way you write. Do not use Claude's LLM machinery to construct sentences. Write in one continuous pass, first sentence to last. Do not draft short and expand. Post it as plain text with no commentary."
)

TURN1_PROMPT = (
    "Read the attached materials carefully. In 6 short bullet points, state only the operative writing constraints you must obey while drafting this chapter. "
    "Do not summarize the plot. Do not discuss themes."
)

TURN2_PROMPT = (
    "Now write the chapter from the outline. Stay inside the source-text habits of mind, sentence logic, and paragraph movement. "
    "Write it straight through in one continuous pass from first sentence to last. Do not draft short and expand. Return plain text only with no commentary."
)


@dataclass
class RunRecord:
    run_id: str
    timestamp: str
    model: str
    temperature: float
    max_tokens: int
    mode: str
    source_file: str
    profiles_file: str
    outline_file: str
    output_file: str
    payload_file: str
    notes: str = ""
    originality_label: str = ""
    originality_score: Optional[float] = None
    manual_notes: str = ""


def load_records() -> pd.DataFrame:
    if CSV_PATH.exists():
        return pd.read_csv(CSV_PATH)
    return pd.DataFrame(
        columns=[
            "run_id",
            "timestamp",
            "model",
            "temperature",
            "max_tokens",
            "mode",
            "source_file",
            "profiles_file",
            "outline_file",
            "output_file",
            "payload_file",
            "notes",
            "originality_label",
            "originality_score",
            "manual_notes",
        ]
    )


def append_record(record: RunRecord) -> None:
    df = load_records()
    df = pd.concat([df, pd.DataFrame([asdict(record)])], ignore_index=True)
    df.to_csv(CSV_PATH, index=False)


def update_record(run_id: str, updates: dict) -> None:
    df = load_records()
    if df.empty:
        return
    mask = df["run_id"] == run_id
    if not mask.any():
        return
    for key, value in updates.items():
        if key in df.columns:
            df.loc[mask, key] = value
    df.to_csv(CSV_PATH, index=False)


def normalize_text(text: str) -> str:
    return (
        text.replace("\u2013", "-")
        .replace("\u2014", "--")
        .replace("\u2018", "'")
        .replace("\u2019", "'")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
        .replace("\u00a0", " ")
    )


def read_docx_bytes(data: bytes) -> str:
    bio = io.BytesIO(data)
    doc = Document(bio)
    paragraphs = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def decode_uploaded_file(uploaded_file) -> Tuple[str, str]:
    suffix = Path(uploaded_file.name).suffix.lower()
    raw = uploaded_file.read()

    if suffix == ".docx":
        text = read_docx_bytes(raw)
    else:
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            text = raw.decode("latin-1")

    return uploaded_file.name, normalize_text(text)



def build_single_turn_payload(source_name: str, source_text: str, profiles_name: str, profiles_text: str, outline_name: str, outline_text: str, user_prompt: str) -> str:
    parts = [
        f"<COMBINED_SOURCE_TEXTS file=\"{source_name}\">\n{source_text.strip()}\n</COMBINED_SOURCE_TEXTS>",
        f"<CHARACTER_PROFILES file=\"{profiles_name}\">\n{profiles_text.strip()}\n</CHARACTER_PROFILES>" if profiles_text.strip() else "",
        f"<OUTLINE file=\"{outline_name}\">\n{outline_text.strip()}\n</OUTLINE>",
        f"<INSTRUCTION>\n{user_prompt.strip()}\n</INSTRUCTION>",
    ]
    return "\n\n".join(p for p in parts if p)



def extract_text_blocks(resp) -> str:
    parts: List[str] = []
    for block in getattr(resp, "content", []) or []:
        if getattr(block, "type", None) == "text":
            parts.append(block.text)
    return "\n".join(parts).strip()



def call_anthropic_single_turn(api_key: str, model: str, system_prompt: str, payload: str, max_tokens: int, temperature: float) -> str:
    if anthropic is None:
        raise RuntimeError("anthropic package is not installed.")
    client = anthropic.Anthropic(api_key=api_key)
    resp = client.messages.create(
        model=model,
        system=system_prompt,
        max_tokens=max_tokens,
        temperature=temperature,
        messages=[{"role": "user", "content": payload}],
    )
    return extract_text_blocks(resp)



def call_anthropic_two_turn(api_key: str, model: str, system_prompt: str, source_name: str, source_text: str, profiles_name: str, profiles_text: str, outline_name: str, outline_text: str, max_tokens: int, temperature: float) -> Tuple[str, str, str]:
    if anthropic is None:
        raise RuntimeError("anthropic package is not installed.")
    client = anthropic.Anthropic(api_key=api_key)

    material_block = build_single_turn_payload(
        source_name=source_name,
        source_text=source_text,
        profiles_name=profiles_name,
        profiles_text=profiles_text,
        outline_name=outline_name,
        outline_text=outline_text,
        user_prompt=TURN1_PROMPT,
    )

    turn1 = client.messages.create(
        model=model,
        system=system_prompt,
        max_tokens=350,
        temperature=0.7,
        messages=[{"role": "user", "content": material_block}],
    )
    anchor_text = extract_text_blocks(turn1)

    turn2_messages = [
        {"role": "user", "content": build_single_turn_payload(source_name, source_text, profiles_name, profiles_text, outline_name, outline_text, TURN1_PROMPT)},
        {"role": "assistant", "content": anchor_text},
        {"role": "user", "content": TURN2_PROMPT},
    ]

    turn2 = client.messages.create(
        model=model,
        system=system_prompt,
        max_tokens=max_tokens,
        temperature=temperature,
        messages=turn2_messages,
    )
    chapter_text = extract_text_blocks(turn2)

    payload_preview = (
        "=== TURN 1 USER ===\n" + material_block + "\n\n"
        "=== TURN 1 ASSISTANT ===\n" + anchor_text + "\n\n"
        "=== TURN 2 USER ===\n" + TURN2_PROMPT
    )

    return anchor_text, payload_preview, chapter_text



def export_zip(df: pd.DataFrame, outputs_root: Path) -> bytes:
    mem = io.BytesIO()
    with zipfile.ZipFile(mem, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("results.csv", df.to_csv(index=False))
        for file_path in sorted(outputs_root.rglob("*")):
            if file_path.is_file():
                zf.write(file_path, arcname=str(file_path.relative_to(outputs_root.parent)))
    mem.seek(0)
    return mem.read()



def main() -> None:
    st.set_page_config(page_title=APP_TITLE, layout="wide")
    st.title(APP_TITLE)
    st.caption("Upload the real source texts, character profiles, and outline. The app saves the exact payload and output for each run.")

    with st.sidebar:
        st.header("Run setup")
        api_key = st.text_input("Anthropic API key", value="", type="password")
        model = st.text_input("Model", value="claude-sonnet-4-6")
        mode = st.selectbox("Run mode", ["single_turn", "two_turn_anchor"], index=1)
        temperature = st.slider("Temperature", 0.0, 1.5, 1.0, 0.1)
        max_tokens = st.number_input("Max output tokens", min_value=300, max_value=12000, value=3000, step=100)
        notes = st.text_input("Batch notes", value="")

    left, right = st.columns([1.15, 0.85])

    with left:
        st.subheader("Materials")
        source_file = st.file_uploader("Combined source texts (.txt, .md, .docx)", type=["txt", "md", "docx"], key="source")
        profiles_file = st.file_uploader("Character profiles (.txt, .md, .docx) - optional", type=["txt", "md", "docx"], key="profiles")
        outline_file = st.file_uploader("Outline (.txt, .md, .docx)", type=["txt", "md", "docx"], key="outline")

        st.subheader("Prompts")
        system_prompt = st.text_area("System prompt", value=DEFAULT_SYSTEM, height=100)
        user_prompt = st.text_area("Primary writing prompt", value=DEFAULT_USER_PROMPT, height=220)

        source_name = ""
        profiles_name = ""
        outline_name = ""
        source_text = ""
        profiles_text = ""
        outline_text = ""

        if source_file is not None:
            source_name, source_text = decode_uploaded_file(source_file)
            st.info(f"Loaded source texts: {source_name}")
            st.code(source_text[:1500], language="text")

        if profiles_file is not None:
            profiles_name, profiles_text = decode_uploaded_file(profiles_file)
            st.info(f"Loaded character profiles: {profiles_name}")
            st.code(profiles_text[:1000], language="text")

        if outline_file is not None:
            outline_name, outline_text = decode_uploaded_file(outline_file)
            st.info(f"Loaded outline: {outline_name}")
            st.code(outline_text[:1200], language="text")

        payload_preview = ""
        if source_text.strip() and outline_text.strip():
            payload_preview = build_single_turn_payload(
                source_name=source_name,
                source_text=source_text,
                profiles_name=profiles_name,
                profiles_text=profiles_text,
                outline_name=outline_name,
                outline_text=outline_text,
                user_prompt=user_prompt,
            )
            st.subheader("Single-turn payload preview")
            st.code(payload_preview[:3500], language="text")

        run_button = st.button("Run generation", type="primary")

        if run_button:
            if not api_key:
                st.error("Enter an Anthropic API key.")
            elif not source_text.strip():
                st.error("Upload the combined source texts file.")
            elif not outline_text.strip():
                st.error("Upload the outline file.")
            elif not user_prompt.strip():
                st.error("Primary writing prompt cannot be empty.")
            else:
                run_id = datetime.now().strftime("%Y%m%d_%H%M%S")
                run_dir = OUTPUTS_DIR / run_id
                run_dir.mkdir(parents=True, exist_ok=True)

                source_path = run_dir / f"{run_id}_source.txt"
                profiles_path = run_dir / f"{run_id}_profiles.txt"
                outline_path = run_dir / f"{run_id}_outline.txt"
                payload_path = run_dir / f"{run_id}_payload.txt"
                output_path = run_dir / f"{run_id}_output.txt"
                meta_path = run_dir / f"{run_id}_meta.json"
                anchor_path = run_dir / f"{run_id}_anchor.txt"

                save_text(source_path, source_text)
                save_text(outline_path, outline_text)
                save_text(profiles_path, profiles_text)

                with st.spinner("Generating..."):
                    try:
                        if mode == "single_turn":
                            payload = build_single_turn_payload(
                                source_name=source_name,
                                source_text=source_text,
                                profiles_name=profiles_name,
                                profiles_text=profiles_text,
                                outline_name=outline_name,
                                outline_text=outline_text,
                                user_prompt=user_prompt,
                            )
                            save_text(payload_path, payload)
                            output_text = call_anthropic_single_turn(
                                api_key=api_key,
                                model=model,
                                system_prompt=system_prompt,
                                payload=payload,
                                max_tokens=int(max_tokens),
                                temperature=float(temperature),
                            )
                        else:
                            anchor_text, payload, output_text = call_anthropic_two_turn(
                                api_key=api_key,
                                model=model,
                                system_prompt=system_prompt,
                                source_name=source_name,
                                source_text=source_text,
                                profiles_name=profiles_name,
                                profiles_text=profiles_text,
                                outline_name=outline_name,
                                outline_text=outline_text,
                                max_tokens=int(max_tokens),
                                temperature=float(temperature),
                            )
                            save_text(anchor_path, anchor_text)
                            save_text(payload_path, payload)

                        save_text(output_path, output_text)

                        meta = {
                            "run_id": run_id,
                            "timestamp": datetime.now().isoformat(timespec="seconds"),
                            "model": model,
                            "temperature": float(temperature),
                            "max_tokens": int(max_tokens),
                            "mode": mode,
                            "source_file": source_name,
                            "profiles_file": profiles_name,
                            "outline_file": outline_name,
                            "source_path": str(source_path),
                            "profiles_path": str(profiles_path),
                            "outline_path": str(outline_path),
                            "payload_path": str(payload_path),
                            "output_path": str(output_path),
                            "notes": notes,
                        }
                        save_text(meta_path, json.dumps(meta, indent=2))

                        append_record(
                            RunRecord(
                                run_id=run_id,
                                timestamp=meta["timestamp"],
                                model=model,
                                temperature=float(temperature),
                                max_tokens=int(max_tokens),
                                mode=mode,
                                source_file=source_name,
                                profiles_file=profiles_name,
                                outline_file=outline_name,
                                output_file=str(output_path),
                                payload_file=str(payload_path),
                                notes=notes,
                            )
                        )

                        st.success("Generation complete.")
                        st.subheader("Output")
                        st.code(output_text, language="text")
                    except Exception as exc:
                        st.error(str(exc))

    with right:
        st.subheader("Run log")
        df = load_records()
        if df.empty:
            st.info("No runs yet.")
        else:
            st.dataframe(df.sort_values("timestamp", ascending=False), use_container_width=True, hide_index=True)

            run_ids = df["run_id"].tolist()
            selected_run = st.selectbox("Select run", run_ids)
            current = df[df["run_id"] == selected_run].iloc[0]

            with st.form("score_form"):
                originality_label = st.text_input("Originality label", value=str(current.get("originality_label", "") or ""))
                originality_score = st.text_input(
                    "Originality score",
                    value="" if pd.isna(current.get("originality_score")) else str(current.get("originality_score")),
                )
                manual_notes = st.text_area("Manual notes", value=str(current.get("manual_notes", "") or ""), height=120)
                submitted = st.form_submit_button("Save")
                if submitted:
                    parsed_score = None
                    raw = originality_score.strip()
                    if raw:
                        try:
                            parsed_score = float(raw)
                        except ValueError:
                            st.error("Originality score must be numeric if provided.")
                            st.stop()
                    update_record(
                        selected_run,
                        {
                            "originality_label": originality_label,
                            "originality_score": parsed_score,
                            "manual_notes": manual_notes,
                        },
                    )
                    st.success("Saved.")
                    st.rerun()

            output_path = Path(str(current["output_file"]))
            payload_path = Path(str(current["payload_file"]))

            if output_path.exists():
                st.markdown("### Selected output")
                st.code(output_path.read_text(encoding="utf-8"), language="text")

            if payload_path.exists():
                st.markdown("### Selected payload")
                st.code(payload_path.read_text(encoding="utf-8")[:3500], language="text")

            zip_bytes = export_zip(df, OUTPUTS_DIR)
            st.download_button(
                "Download outputs + CSV",
                data=zip_bytes,
                file_name="author_runs_export.zip",
                mime="application/zip",
            )

    st.markdown("---")
    st.write(
        "Recommended first test: single_turn, claude-sonnet-4-6, temperature 1.0, max output tokens 3000, using the exact source texts, profiles, and outline you want Claude to read."
    )


if __name__ == "__main__":
    main()
